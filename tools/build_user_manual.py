from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


APP_DIR = Path(__file__).resolve().parents[1]
OUT = APP_DIR / "docs" / "World3-03_Shiny_App_User_Manual.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    for paragraph in cell.paragraphs:
      paragraph.paragraph_format.space_after = Pt(0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True)
        set_cell_shading(hdr[i], "E8EEF5")
        if widths:
            hdr[i].width = widths[i]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
            if widths:
                cells[i].width = widths[i]
    doc.add_paragraph()
    return table


def style_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color in [
        ("Heading 1", 16, RGBColor(46, 116, 181)),
        ("Heading 2", 13, RGBColor(46, 116, 181)),
        ("Heading 3", 12, RGBColor(31, 77, 120)),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True


def build():
    doc = Document()
    style_doc(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("World3-03 Shiny App User Manual")
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(31, 77, 120)

    subtitle = doc.add_paragraph()
    subtitle.add_run("Guide to scenario controls, animation mode, and model outputs").italic = True

    doc.add_paragraph(
        "This manual explains how to use the World3-03 Simulation Cockpit. "
        "The app runs a native R port of the local PyWorld3-03 model and lets "
        "you explore how policy timing and selected assumptions affect population, "
        "resources, food, industrial output, pollution, human welfare, and footprint indicators."
    )

    doc.add_heading("Quick Start", level=1)
    add_bullets(doc, [
        "Open the RWorld3-03 project in RStudio.",
        "Run shiny::runApp(\"app\") from the project root, or shiny::runApp() from inside the app folder.",
        "Choose a scenario preset or adjust the sliders.",
        "Click Run simulation after changing assumptions.",
        "Use Play, Pause, Reset, Current year, and Animation speed to reveal the run year by year.",
    ])

    doc.add_heading("What the App Is Doing", level=1)
    doc.add_paragraph(
        "Each run computes the full World3-03 trajectory from 1900 through 2100. "
        "The presentation timeline does not change the model; it controls how much "
        "of the already-computed trajectory is visible. This makes live presentation smoother "
        "while preserving the exact simulation result for the selected assumptions."
    )

    doc.add_heading("Splash Screen and Icon", level=1)
    doc.add_paragraph(
        "The app opens with a dark splash screen titled MIT World3-03 Model. "
        "The background uses stylized, faded model-trajectory lines to set the tone for a "
        "systems-collapse scenario tool without using any official MIT seal or trademarked logo. "
        "The custom app icon combines an earth form with a trajectory curve, signaling that the "
        "application is an exploratory interface for the World3-03 model rather than an official MIT product."
    )

    doc.add_heading("Scenario Presets", level=1)
    add_table(
        doc,
        ["Preset", "What it means"],
        [
            ["Default Python run", "Uses the PyWorld3-03 defaults: general policy year 1975, with resource technology, pollution technology, yield technology, and resource-allocation policy timing left inactive at year 4000."],
            ["2004 reference run", "Defers the broad general policy switch to 4000. This is closer to the reference-style run in the included Python scenario script."],
            ["More resources", "Doubles initial nonrenewable resources while keeping major policy and technology switches inactive."],
            ["More resources + pollution control", "Doubles initial nonrenewable resources and starts pollution technology in 2002."],
        ],
        widths=[Inches(2.0), Inches(4.4)],
    )

    doc.add_heading("Important Convention: Year 4000", level=1)
    doc.add_paragraph(
        "A policy year of 4000 is effectively an off switch for a model run that ends in 2100. "
        "It does not mean the app is literally simulating to year 4000. It means the policy, "
        "technology, or allocation change is delayed beyond the simulation horizon and therefore "
        "does not activate during the visible run."
    )

    doc.add_heading("Controls", level=1)
    add_table(
        doc,
        ["Control", "Type", "Interpretation"],
        [
            ["Scenario preset", "Preset selector", "Loads a coherent starting bundle of policy years and assumption values. You can still edit the sliders after choosing a preset."],
            ["General policy year", "Timing control", "Switches model relationships that use the general policy year. Earlier years mean policy-related table changes activate sooner; 4000 keeps them inactive."],
            ["Pollution technology year", "Timing control", "Year when pollution-control technology begins changing the persistent pollution generation factor. 4000 means no pollution technology shift during 1900-2100."],
            ["Resource technology year", "Timing control", "Year when resource-use technology begins changing the nonrenewable resource usage factor. 4000 means no resource technology shift during the run."],
            ["Agricultural yield technology year", "Timing control", "Year when yield technology begins responding to food-ratio pressure. 4000 means no yield technology shift during the run."],
            ["Resource extraction allocation year", "Timing control", "Year when the model can switch to the alternate relationship for the fraction of capital allocated to obtaining resources."],
            ["Initial nonrenewable resources", "Magnitude control", "Multiplier on the starting nonrenewable resource stock. A value of 1x is the baseline; 2x doubles the starting resource stock."],
            ["Desired food ratio", "Magnitude/control target", "Food adequacy target used in the yield-technology response. Larger values represent a higher desired food buffer relative to subsistence food per capita."],
            ["Run simulation", "Command", "Re-runs the model using the current control values."],
            ["Play / Pause / Reset", "Presentation controls", "Reveal or stop the trajectory over time. Reset returns the visible year to 1900."],
            ["Current year", "Presentation control", "Shows results only through the selected year. Drag manually or use Play."],
            ["Animation speed", "Presentation control", "Controls the delay between visible years. Smaller values are faster; larger values are slower."],
        ],
        widths=[Inches(1.65), Inches(1.35), Inches(3.45)],
    )

    doc.add_heading("Outputs", level=1)
    add_table(
        doc,
        ["Output", "What to look for"],
        [
            ["Metric cards", "Peak population, resources remaining at the visible year, peak pollution so far, and human welfare at the visible year."],
            ["Overview tab", "Scaled headline indicators on one chart: population, nonrenewable resources, food, industrial output, pollution, and human welfare."],
            ["Population tab", "Total population and age cohorts. Useful for seeing demographic momentum and collapse/recovery patterns."],
            ["Capital tab", "Industrial output per capita, service output per capita, and consumption of industrial output per capita."],
            ["Agriculture tab", "Food per capita, land yield, and land fertility. Useful for diagnosing food-system stress."],
            ["Pollution tab", "Persistent pollution index, pollution-generation factor, pollution technology, and ecological footprint."],
            ["Resources tab", "Nonrenewable resource fraction remaining, capital allocated to resource extraction, resource usage factor, and resource technology."],
            ["Data tab", "The last visible rows of the headline data table for inspection or copying."],
        ],
        widths=[Inches(1.6), Inches(4.85)],
    )

    doc.add_heading("How to Present a Scenario", level=1)
    add_bullets(doc, [
        "Start with the Default Python run or 2004 reference run preset.",
        "Click Run simulation, then click Reset so the visible year is 1900.",
        "Open the Overview tab and click Play.",
        "Pause around turning points, such as peak population, rapid pollution growth, resource decline, or falling food per capita.",
        "Switch to a sector tab to explain the mechanism behind the visible change.",
        "Change one assumption at a time, click Run simulation, and replay the same years for comparison.",
    ])

    doc.add_heading("Reading the Results Carefully", level=1)
    doc.add_paragraph(
        "The app is designed for scenario exploration rather than prediction. "
        "A control change should be read as: under this model structure and this assumption set, "
        "the system behaves this way. The most useful comparisons usually come from changing one "
        "lever at a time and watching which sector moves first."
    )

    doc.add_heading("Notes for Publishing", level=1)
    add_bullets(doc, [
        "The publishable app root is the app folder.",
        "The app is self-contained: it includes app.R, R/world3_03.R, data/functions_table_world3.json, www/app.css, and www/logo.svg.",
        "Generate the Posit manifest from the app folder with source(\"tools/write_manifest.R\").",
        "If publishing from GitHub, use the app folder as the Shiny app root.",
    ])

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
